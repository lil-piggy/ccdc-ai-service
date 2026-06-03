from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_heading_style(run, size=16, bold=True, color=RGBColor(0x1a, 0x1a, 0x1a)):
    font = run.font
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = color
    font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def set_body_style(run, size=11):
    font = run.font
    font.size = Pt(size)
    font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    size = {1: 20, 2: 16, 3: 14}.get(level, 12)
    set_heading_style(run, size=size, bold=True, color=RGBColor(0x1a, 0x1a, 0x1a))
    return p

def add_para(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_body_style(run)
    if bold:
        run.font.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_body_style(run)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_body_style(run)
    return p

doc = Document()

# 封面
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CCDC 智能金融助手\n操作手册')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('\nv1.0 | 2026-05-28')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# 目录标题
add_heading_custom(doc, '目录', level=1)
add_para(doc, '一、系统概述', indent=False)
add_para(doc, '二、快速开始', indent=False)
add_para(doc, '三、用户认证', indent=False)
add_para(doc, '四、核心功能', indent=False)
add_para(doc, '五、快捷键大全', indent=False)
add_para(doc, '六、管理配置', indent=False)
add_para(doc, '七、常见问题', indent=False)
doc.add_page_break()

# 一、系统概述
add_heading_custom(doc, '一、系统概述', level=1)
add_para(doc, 'CCDC 智能金融助手（代号：小发）是一款面向债券市场从业者的 AI 智能客服系统，基于大语言模型技术，提供债券定价、信用分析、募集书解读、宏观点评等专业金融服务。')

add_heading_custom(doc, '1.1 系统架构', level=2)
add_bullet(doc, '前端：单页应用（SPA），HUD 科技风格界面，支持三套主题')
add_bullet(doc, '后端：Node.js + Express，JWT 认证，多用户数据隔离')
add_bullet(doc, '数据库：SQLite（sql.js），文件持久化，支持用户/对话/配置/知识库四张表')
add_bullet(doc, 'AI 引擎：OpenAI 兼容接口代理，支持流式 SSE 输出')

add_heading_custom(doc, '1.2 部署信息', level=2)
add_bullet(doc, '本地地址：http://localhost:3001')
add_bullet(doc, '线上地址：https://ccdc-ai-service.onrender.com')
add_bullet(doc, '默认模型：Kimi K2.6（由管理员统一配置）')

doc.add_page_break()

# 二、快速开始
add_heading_custom(doc, '二、快速开始', level=1)

add_heading_custom(doc, '2.1 首次使用', level=2)
add_number(doc, '打开浏览器，访问服务地址')
add_number(doc, '点击「注册」按钮，输入用户名（至少3位）和密码（至少6位）')
add_number(doc, '注册成功后自动登录，进入主界面')
add_number(doc, '在底部输入框输入问题，按回车或点击发送按钮开始对话')

add_heading_custom(doc, '2.2 界面布局', level=2)
add_para(doc, '主界面分为三个区域：')
add_bullet(doc, '左侧边栏：品牌 Logo、系统状态、新会话按钮、历史对话列表、快捷模板')
add_bullet(doc, '中间主区域：对话展示区，支持 Markdown 渲染、代码高亮')
add_bullet(doc, '底部输入区：文本输入框、发送按钮、附件上传')

doc.add_page_break()

# 三、用户认证
add_heading_custom(doc, '三、用户认证', level=1)

add_heading_custom(doc, '3.1 注册账号', level=2)
add_para(doc, '首次使用需注册账号。系统采用自研用户名密码体系，非 GitHub/邮箱登录。')
add_para(doc, '限制条件：')
add_bullet(doc, '用户名：至少 3 个字符，全局唯一')
add_bullet(doc, '密码：至少 6 个字符')
add_bullet(doc, '注册成功后自动获得 7 天有效期的 JWT Token')

add_heading_custom(doc, '3.2 登录与退出', level=2)
add_para(doc, '已注册用户点击「登录」按钮，输入用户名和密码即可进入。')
add_para(doc, '点击用户头像或侧边栏底部的「退出登录」按钮可安全退出，本地 Token 会被清除。')

add_heading_custom(doc, '3.3 Token 机制', level=2)
add_para(doc, '系统使用 JWT（JSON Web Token）进行身份认证。Token 有效期为 7 天，过期后需要重新登录。所有 API 请求（除注册/登录外）都需在请求头中携带 Authorization: Bearer <token>。')

doc.add_page_break()

# 四、核心功能
add_heading_custom(doc, '四、核心功能', level=1)

add_heading_custom(doc, '4.1 新会话（开启新分析会话）', level=2)
add_para(doc, '点击左侧边栏顶部的「+ 开启新分析会话」按钮，或按 Ctrl+N 快捷键，可创建一个新的空白对话。')
add_para(doc, '每个会话独立保存，支持随时切换和删除。')

add_heading_custom(doc, '4.2 历史对话管理', level=2)
add_para(doc, '左侧边栏展示当前用户的所有历史对话，按更新时间倒序排列。支持以下操作：')
add_bullet(doc, '点击对话标题切换会话')
add_bullet(doc, '鼠标悬停显示删除按钮，点击可删除单条对话')
add_bullet(doc, '支持置顶重要对话（pinned 标记）')
add_bullet(doc, '对话数据自动同步到服务端，换设备登录后历史记录不丢失')

add_heading_custom(doc, '4.3 快捷模板', level=2)
add_para(doc, '左侧边栏「快捷模板」区域提供 5 个一键提示词模板：')
add_bullet(doc, '债券定价：快速分析债券的估值、利差、定价合理性')
add_bullet(doc, '信用分析：评估发行人信用资质、财务指标、风险等级')
add_bullet(doc, '募集书解读：自动提取募集说明书核心条款、风险提示')
add_bullet(doc, '宏观点评：基于最新宏观数据生成债券市场点评')
add_bullet(doc, '收益率曲线：分析收益率曲线形态、期限利差变化')
add_para(doc, '点击任意模板，系统会自动在输入框填入预设提示词，用户可在此基础上修改后发送。')

add_heading_custom(doc, '4.4 知识库（F11）', level=2)
add_para(doc, '知识库功能允许用户上传自定义文档，AI 在回答时会自动检索相关知识库内容作为参考（RAG 检索增强生成）。')
add_para(doc, '操作步骤：')
add_number(doc, '按 F11 或点击「知识库」按钮打开面板')
add_number(doc, '点击「上传文件」，选择 txt/pdf/doc 等文本文件')
add_number(doc, '上传成功后，文档会出现在知识库列表中')
add_number(doc, '发送消息时，系统会自动匹配知识库中与问题相关的片段，附加到提示词中')
add_para(doc, '注意事项：')
add_bullet(doc, '知识库文档按用户隔离，其他用户无法查看你的知识库')
add_bullet(doc, '支持上传多个文档，系统会分别检索每个文档的片段')
add_bullet(doc, '可点击文档旁的删除按钮移除不需要的文档')

add_heading_custom(doc, '4.5 角色定制 / Persona（F12）', level=2)
add_para(doc, '通过 Persona 功能，可以自定义 AI 的角色设定，让回答更符合特定场景需求。')
add_para(doc, '预设角色包括：')
add_bullet(doc, '债券分析师：专注于债券定价、利差分析')
add_bullet(doc, '信用评级师：专注于信用风险评估')
add_bullet(doc, '宏观研究员：专注于宏观经济与政策解读')
add_para(doc, '用户也可以自定义新的 Persona，设置角色名称、系统提示词和描述。')

add_heading_custom(doc, '4.6 多模型对比（F16）', level=2)
add_para(doc, '开启多模型对比模式后，同一个问题会同时发送给多个模型，并排展示不同模型的回答，便于对比分析。')
add_para(doc, '注意：此功能需要管理员在后台配置多个模型，当前部署默认只配置了一个模型。')

add_heading_custom(doc, '4.7 自动标题生成（F19）', level=2)
add_para(doc, '系统会根据对话内容自动生成简洁的标题，便于在历史列表中快速识别对话主题。')
add_para(doc, '按 F19 可手动触发重新生成当前对话的标题。')

add_heading_custom(doc, '4.8 加密存储（F21）', level=2)
add_para(doc, '开启加密存储后，本地缓存的对话历史会经过加密处理，提升数据安全性。')
add_para(doc, '加密后的数据以 finChatVault_v3 为 key 存储在 localStorage 中。')

add_heading_custom(doc, '4.9 备份与恢复（F22）', level=2)
add_para(doc, '支持将对话历史、配置、知识库等数据导出为加密文件，用于备份或迁移。')
add_para(doc, '导出后，在新设备或新浏览器中可通过「恢复」功能导入备份文件，还原全部数据。')

add_heading_custom(doc, '4.10 统计面板（F23）', level=2)
add_para(doc, '统计面板展示用户的使用数据，包括：')
add_bullet(doc, '累计对话次数')
add_bullet(doc, '累计发送消息数')
add_bullet(doc, '累计接收 Token 数')
add_bullet(doc, '使用天数统计')
add_para(doc, '数据按用户隔离，仅展示当前登录用户的数据。')

add_heading_custom(doc, '4.11 主题切换（F4）', level=2)
add_para(doc, '系统内置三套主题风格：')
add_bullet(doc, '深色主题（Dark）：默认主题，深蓝黑色调，适合夜间使用')
add_bullet(doc, '浅色主题（Light）：白色背景，适合日间使用')
add_bullet(doc, '赛博主题（Cyber）：霓虹紫青色调，科技感最强')
add_para(doc, '按 F4 键可快速切换主题，或点击顶部工具栏的主题按钮。')

add_heading_custom(doc, '4.12 实时市场数据（F6）', level=2)
add_para(doc, '按 F6 可打开实时市场数据面板，展示债券市场关键指标（模拟数据）。')
add_para(doc, '包括：中债收益率曲线、信用利差、成交量等实时行情。')

add_heading_custom(doc, '4.13 导出功能', level=2)
add_para(doc, '点击顶部工具栏的「导出」按钮，可将当前对话内容导出为 Markdown 或文本文件，便于保存和分享。')

add_heading_custom(doc, '4.14 API 配置查看', level=2)
add_para(doc, '点击「API 配置」按钮可查看当前使用的模型信息。由于采用管理员统一配置模式，普通用户无法修改 API Key 和接口地址，只能查看当前生效的模型名称。')

doc.add_page_break()

# 五、快捷键大全
add_heading_custom(doc, '五、快捷键大全', level=1)
add_para(doc, '以下快捷键在聊天界面全局生效：')

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '快捷键'
hdr_cells[1].text = '功能'
hdr_cells[2].text = '说明'

shortcuts = [
    ('Ctrl + N', '新建会话', '开启一个新的空白对话'),
    ('F4', '切换主题', '在 Dark / Light / Cyber 三套主题间循环切换'),
    ('F6', '实时市场数据', '打开/关闭实时行情面板'),
    ('F11', '知识库', '打开/关闭知识库管理面板'),
    ('F12', '角色定制', '打开/关闭 Persona 设置面板'),
    ('F16', '多模型对比', '开启/关闭多模型并排对比模式'),
    ('F19', '自动标题', '手动触发重新生成当前对话标题'),
    ('F21', '加密存储', '开启/关闭本地数据加密'),
    ('F22', '备份恢复', '打开数据导出/导入面板'),
    ('F23', '统计面板', '打开用户使用数据统计'),
    ('Enter', '发送消息', '在输入框中按回车发送消息'),
    ('Shift + Enter', '换行', '输入框中换行而不发送'),
]

for key, func, desc in shortcuts:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = func
    row_cells[2].text = desc
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_body_style(run)

doc.add_page_break()

# 六、管理配置
add_heading_custom(doc, '六、管理配置', level=1)

add_heading_custom(doc, '6.1 环境变量（Render 部署）', level=2)
add_para(doc, '管理员通过 Render Dashboard 的 Environment 页面配置以下变量：')

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = '变量名'
hdr_cells[1].text = '必填'
hdr_cells[2].text = '说明'

env_vars = [
    ('ADMIN_API_URL', '是', 'OpenAI 兼容的 API 接口地址'),
    ('ADMIN_API_KEY', '是', 'API 接口的认证密钥'),
    ('ADMIN_API_MODEL', '是', '默认使用的模型名称，如 kimi-k2.6'),
    ('JWT_SECRET', '是', 'JWT 签名的密钥，建议 32 位以上随机字符串'),
    ('PORT', '否', '服务端口，默认 3001'),
]

for name, required, desc in env_vars:
    row_cells = table2.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = required
    row_cells[2].text = desc
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_body_style(run)

add_heading_custom(doc, '6.2 数据库持久化', level=2)
add_para(doc, '系统使用 sql.js（纯 JavaScript SQLite）作为数据库，数据以文件形式（data.sqlite）持久化存储在磁盘上。')
add_para(doc, '每次写入操作后，数据库会自动导出为文件；服务启动时，会从文件恢复数据库。这意味着即使服务重启或部署更新，用户数据也不会丢失。')

add_heading_custom(doc, '6.3 用户数据隔离', level=2)
add_para(doc, '所有用户数据严格按 user_id 隔离：')
add_bullet(doc, '对话历史：每个用户只能查看和管理自己的对话')
add_bullet(doc, '用户配置：主题、模型偏好等按用户独立存储')
add_bullet(doc, '知识库：每个用户的知识库文档互不可见')

doc.add_page_break()

# 七、常见问题
add_heading_custom(doc, '七、常见问题', level=1)

add_heading_custom(doc, 'Q1：为什么聊天返回「上游 API 错误：Invalid Authentication」？', level=2)
add_para(doc, 'A：说明管理员配置的 ADMIN_API_KEY 无效或已过期。请联系管理员检查 Render 环境变量中的 API Key 是否正确。')

add_heading_custom(doc, 'Q2：历史对话在换电脑后还能看见吗？', level=2)
add_para(doc, 'A：可以。所有对话数据存储在服务端数据库中，只要用同一账号登录，任何设备都能访问完整的历史记录。')

add_heading_custom(doc, 'Q3：知识库上传的文件有大小限制吗？', level=2)
add_para(doc, 'A：当前没有严格的大小限制，但建议单个文件不超过 5MB，以保证检索效率。')

add_heading_custom(doc, 'Q4：服务休眠后数据会丢失吗？', level=2)
add_para(doc, 'A：不会。Render 免费版实例在 15 分钟无访问后会休眠，但数据库文件（data.sqlite）保存在磁盘上，唤醒后数据完整恢复。')

add_heading_custom(doc, 'Q5：如何修改默认模型？', level=2)
add_para(doc, 'A：普通用户无法修改模型。只有管理员可以在 Render 环境变量中修改 ADMIN_API_MODEL 的值，修改后需要重新部署生效。')

add_heading_custom(doc, 'Q6：Token 过期了怎么办？', level=2)
add_para(doc, 'A：Token 有效期为 7 天。过期后系统会自动弹出登录框，重新输入用户名密码登录即可获取新 Token。')

add_heading_custom(doc, 'Q7：免费版 Render 有什么限制？', level=2)
add_para(doc, 'A：Render 免费版的限制包括：')
add_bullet(doc, '15 分钟无访问后实例休眠，首次访问需等待 30 秒唤醒')
add_bullet(doc, '每月 750 小时运行时间上限')
add_bullet(doc, '不支持 SSH 访问和持久化磁盘（但本项目的数据库文件自动保存）')
add_bullet(doc, '带宽和性能有限，适合个人和小团队使用')

# 保存
output_path = r'D:\桌面\TEMP\USB_DATA\ccdc-ai-service\CCDC_Manual.docx'
doc.save(output_path)
print('文档已生成：' + output_path)
